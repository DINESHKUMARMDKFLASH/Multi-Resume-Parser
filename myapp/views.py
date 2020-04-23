from django.shortcuts import render, redirect
from docx import *
from .forms import StudentsForm
from .models import Students

# Create your views here.

def index(request):
    if request.method == 'POST':
        try:
            if request.FILES['document']:
                document = Document(request.FILES['document'])

                lines = []
                for para in document.paragraphs:
                    temp = ""
                    line = (para.text).split()
                    c = 0
                    for x in line:
                        if (x == ":"):
                            c = 1
                        if (c > 1):
                            temp += x + " "
                        c += 1
                    lines.append(temp)

                lcontext = {'fname': lines[0],
                            'lname': lines[1],
                            'mail': lines[2],
                            'phno': lines[3],
                            'st': lines[4],
                            'city': lines[5],
                            'state': lines[6],
                            'country': lines[7],
                            'pincode': lines[8],
                            'work': lines[9],
                            'edu': lines[10],
                            'skill': lines[11],
                            'workexp': lines[12]
                            }

                form = StudentsForm(lcontext)
                if form.is_valid():
                    try:
                        form.save()
                        return redirect('/view')
                    except:
                        return redirect('/view')
                else:
                    return redirect('/index')
        except:
            return redirect('/view')

    return render(request, 'index.html')


def view(request):
    students = Students.objects.all()
    return render(request, "view.html", {'students': students})


def resume(request):
    students = Students.objects.all()
    return render(request, "resume.html", {'students': students})


def addmore(request):
    if request.method == 'POST':
        try:
            if request.FILES['document']:
                document = Document(request.FILES['document'])

                lines = []
                for para in document.paragraphs:
                    temp = ""
                    line = (para.text).split()
                    c = 0
                    for x in line:
                        if (x == ":"):
                            c = 1
                        if (c > 1):
                            temp += x + " "
                        c += 1
                    lines.append(temp)

                lcontext = {'fname': lines[0],
                            'lname': lines[1],
                            'mail': lines[2],
                            'phno': lines[3],
                            'st': lines[4],
                            'city': lines[5],
                            'state': lines[6],
                            'country': lines[7],
                            'pincode': lines[8],
                            'work': lines[9],
                            'edu': lines[10],
                            'skill': lines[11],
                            'workexp': lines[12]
                            }

                form = StudentsForm(lcontext)
                if form.is_valid():
                    try:
                        form.save()
                        return redirect('/view')
                    except:
                        pass
                else:
                    return redirect('/index')
        except:
            return redirect('/view')

    return render(request, 'index.html')


def noofres(request):
    nooflist = {}
    if request.method == 'POST':
        no = request.POST['resume']
        for i in range(int(no)):
            nooflist[i] = i

        return render(request, 'index.html')

    return render(request, 'noofres.html')

def welcome(request):


    return render(request,'welcome.html')

def delete(request,id):
    students=Students.objects.get(id=id)
    students.delete()
    return redirect("/view")

def curview(request,id):
    students = Students.objects.get(id=id)
    return render(request, "curview.html", {'students': students})

def edit(request,id):
    students = Students.objects.get(id=id)
    return render(request, "edit.html", {'students': students})

def update(request,id):
    if request.method == 'POST':
        students = Students.objects.get(id=id)
        students.delete()
        form = StudentsForm(request.POST)
        print(form)
        if form.is_valid():
            try:
                form.save()
                return redirect('/view')
            except:
                pass
        else:
            form = StudentsForm()